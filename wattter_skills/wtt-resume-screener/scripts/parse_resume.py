#!/usr/bin/env python3
"""简历解析器：支持 PDF / DOCX / Markdown / TXT。

提取候选人的结构化信息：基础信息、教育背景、工作经历、项目经历、技能清单。
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any


# ============ 文本提取 ============

def extract_text_from_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        from PyPDF2 import PdfReader  # type: ignore
    reader = PdfReader(str(path))
    parts = [p.extract_text() or "" for p in reader.pages]
    return "\n".join(parts)


def extract_text_from_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_md(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix in {".docx", ".doc"}:
        return extract_text_from_docx(path)
    if suffix in {".md", ".markdown"}:
        return extract_text_from_md(path)
    if suffix in {".txt", ""}:
        return path.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported file type: {suffix}")


# ============ 结构化字段抽取 ============

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[- ]?)?(?:\(?\d{2,4}\)?[- ]?)?\d{3,4}[- ]?\d{4}")
YEAR_RE = re.compile(r"(19|20)\d{2}")
SCHOOL_HINTS = ("大学", "学院", "University", "Institute", "School", "College")
COMPANY_HINTS = (
    "公司", "科技", "网络", "信息", "集团", "有限", "工作室", "Inc", "Ltd", "Corp",
    "公司", "Co.,", "LLC",
)


def _first_non_empty_line(text: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if line and not line.startswith("#"):
            return line[:80]
    return ""


def guess_name(text: str) -> str:
    """从首部 10 行猜姓名：通常 2-4 个汉字、无空格、不含特殊符号。"""
    for line in text.splitlines()[:10]:
        line = line.strip()
        if not line:
            continue
        if "@" in line or any(k in line for k in ("简历", "Resume", "CV")):
            continue
        # 纯中文 2-4 字
        if re.fullmatch(r"[一-龥]{2,4}", line):
            return line
        # 英文 First Last
        if re.fullmatch(r"[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}", line):
            return line
    return ""


def extract_basics(text: str) -> dict[str, str]:
    email_match = EMAIL_RE.search(text)
    phone_match = PHONE_RE.search(text)
    return {
        "name": guess_name(text),
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0) if phone_match else "",
    }


def split_sections(text: str) -> dict[str, str]:
    """按常见小标题切分简历段落。"""
    headers = [
        "教育经历", "教育背景", "Education", "工作经历", "工作经验",
        "Work Experience", "项目经历", "项目经验", "Projects",
        "技能", "专业技能", "Skills", "获奖", "荣誉", "Awards",
        "自我评价", "个人评价", "Summary", "About",
    ]
    pattern = re.compile(
        r"^[\s#>*\-]*(" + "|".join(re.escape(h) for h in headers) + r")[\s:：]*$",
        re.IGNORECASE | re.MULTILINE,
    )
    matches = list(pattern.finditer(text))
    sections: dict[str, str] = {}
    if not matches:
        sections["_raw"] = text
        return sections
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        key = m.group(1).strip().lower()
        sections[key] = text[start:end].strip()
    return sections


def extract_skills(text: str) -> list[str]:
    """从全文 / 技能段落中粗略提取技能关键词。"""
    skill_keywords = {
        # 编程语言
        "Python", "Java", "Go", "Golang", "Rust", "C++", "C#", "C ", "JavaScript",
        "TypeScript", "Ruby", "PHP", "Swift", "Kotlin", "Scala", "R ", "MATLAB",
        "SQL", "Shell", "Bash",
        # 前端
        "React", "Vue", "Angular", "Svelte", "Next.js", "Nuxt", "Webpack", "Vite",
        "Tailwind", "CSS", "HTML",
        # 后端 / 框架
        "Spring", "SpringBoot", "Django", "Flask", "FastAPI", "Node.js", "Express",
        "NestJS", "Rails", "Gin", "Echo",
        # 数据 / AI
        "TensorFlow", "PyTorch", "Paddle", "HuggingFace", "LangChain", "LangGraph",
        "Scikit-learn", "Pandas", "NumPy", "Spark", "Hadoop", "Flink", "Kafka",
        "Airflow", "dbt", "ClickHouse", "Redis", "MongoDB", "PostgreSQL", "MySQL",
        "ElasticSearch", "Snowflake",
        # 云原生 / DevOps
        "Docker", "Kubernetes", "K8s", "Helm", "Terraform", "Ansible", "Jenkins",
        "GitLab CI", "GitHub Actions", "AWS", "Azure", "GCP", "阿里云", "腾讯云",
        # 移动端
        "iOS", "Android", "Flutter", "React Native", "Weex",
        # 设计 / 产品
        "Figma", "Sketch", "Axure", "Photoshop", "Illustrator", "Premiere",
        "After Effects",
        # 其他
        "Git", "Linux", "Agile", "Scrum",
    }
    found: list[str] = []
    for kw in skill_keywords:
        # 用单词边界做粗略匹配，避免误报
        pattern = re.escape(kw.strip())
        if re.search(rf"(?<![A-Za-z]){pattern}(?![A-Za-z])", text):
            found.append(kw.strip())
    return found


def extract_education(sections: dict[str, str], full_text: str) -> list[dict[str, str]]:
    edu_text = sections.get("教育经历") or sections.get("教育背景") or sections.get("education") or ""
    if not edu_text:
        for line in full_text.splitlines():
            if any(h in line for h in SCHOOL_HINTS):
                edu_text = line
                break
    items: list[dict[str, str]] = []
    if not edu_text:
        return items
    for line in edu_text.splitlines():
        line = line.strip()
        if not line:
            continue
        years = re.findall(r"(?:19|20)\d{2}", line)
        period = ""
        if years:
            uniq = sorted(set(years))
            period = uniq[0] + ("-" + uniq[-1] if len(uniq) > 1 else "")
        school = next((h for h in SCHOOL_HINTS if h in line), "")
        degree = ""
        for d in ("博士", "硕士", "本科", "学士", "大专"):
            if d in line:
                degree = d
                break
        items.append({"period": period, "school": school, "degree": degree, "raw": line})
    return items


def extract_experience(sections: dict[str, str], full_text: str) -> list[dict[str, str]]:
    """工作/项目经历统一抽取。"""
    exp_text = (
        sections.get("工作经历") or sections.get("工作经验") or sections.get("work experience")
        or sections.get("项目经历") or sections.get("项目经验") or sections.get("projects")
        or ""
    )
    items: list[dict[str, str]] = []
    if not exp_text:
        return items
    blocks = re.split(r"\n\s*\n", exp_text)
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        first_line = block.splitlines()[0]
        period = ""
        years = re.findall(r"(?:19|20)\d{2}", first_line)
        if years:
            uniq = sorted(set(years))
            period = uniq[0] + ("-" + uniq[-1] if len(uniq) > 1 else "")
        items.append({"period": period, "raw": block})
    return items


# ============ 主流程 ============

def parse_resume(path: Path) -> dict[str, Any]:
    text = extract_text(path)
    sections = split_sections(text)
    basics = extract_basics(text)
    return {
        "source_file": str(path),
        "char_count": len(text),
        "basics": basics,
        "sections": list(sections.keys()),
        "education": extract_education(sections, text),
        "experience": extract_experience(sections, text),
        "skills": extract_skills(text),
        "text": text,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="简历解析器：抽取基础信息/教育/经历/技能")
    parser.add_argument("input", nargs="+", help="一个或多个简历文件（PDF/DOCX/MD/TXT）")
    parser.add_argument("--output", "-o", default="-", help="输出 JSON 文件路径，默认 stdout")
    args = parser.parse_args()

    results: list[dict[str, Any]] = []
    for src in args.input:
        p = Path(src)
        if not p.exists():
            print(f"[WARN] file not found: {src}", file=sys.stderr)
            continue
        try:
            results.append(parse_resume(p))
        except Exception as e:  # noqa: BLE001
            print(f"[ERROR] failed to parse {src}: {e}", file=sys.stderr)
            results.append({"source_file": str(p), "error": str(e)})

    payload = json.dumps(results, ensure_ascii=False, indent=2)
    if args.output == "-":
        print(payload)
    else:
        Path(args.output).write_text(payload, encoding="utf-8")
        print(f"[OK] {len(results)} resumes parsed -> {args.output}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
